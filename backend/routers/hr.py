"""
routers/hr.py — Endpoints RH complets
PE, Autorisation Sortie, Reprise, Maladie, Fin Manquant, RC
+ date_creation manuelle sur tous les formulaires
"""

import io
import re
import os
import tempfile
from datetime import datetime, date
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel

import models, auth
from database import get_db

router = APIRouter(prefix="/hr", tags=["RH"])
TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"


# =============================================================================
# HELPERS
# =============================================================================

def split_jma(d: date):
    return f"{d.day:02d}", f"{d.month:02d}", f"{d.year:04d}"

def replace_placeholders(doc, mapping: dict):
    def _replace_para(paragraph):
        for key, val in mapping.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(val))
    for p in doc.paragraphs:
        _replace_para(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_para(p)

def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    import pythoncom
    from docx2pdf import convert
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "doc.docx")
        pdf_path  = os.path.join(tmp, "doc.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
        finally:
            try: pythoncom.CoUninitialize()
            except: pass
        with open(pdf_path, "rb") as f:
            return f.read()

def stream_pdf(pdf_bytes: bytes, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

def save_doc_record(db, user_id: int, type_doc: str, metadata: dict):
    record = models.DocumentRH(user_id=user_id, type_doc=type_doc, metadata_doc=metadata)
    db.add(record); db.commit(); db.refresh(record)
    return record

def fmt_heure(h: str) -> str:
    if ":" not in h: return h
    hh, mm = h.split(":", 1)
    return f"{hh.zfill(2)}h : {mm.zfill(2)} mn"

def get_original_user(db, record, current_user):
    if current_user.role != "admin" and record.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Accès refusé")
    user = db.query(models.User).filter(models.User.id == record.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Agent introuvable")
    return user

def parse_date_creation(date_creation: Optional[str]) -> str:
    """
    Parse date_creation from frontend (YYYY-MM-DD) to display format (DD/MM/YYYY).
    Falls back to today if None or invalid.
    """
    if date_creation:
        try:
            return date.fromisoformat(date_creation).strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            pass
    return date.today().strftime("%d/%m/%Y")


# =============================================================================
# SCHEMAS — date_creation ajouté sur tous les formulaires
# =============================================================================

class PERequest(BaseModel):
    type_perm:     str
    date1:         Optional[str] = None
    date2:         Optional[str] = None
    heure_debut:   Optional[str] = None
    heure_fin:     Optional[str] = None
    date_creation: Optional[str] = None   # ← NOUVEAU

class SortieRequest(BaseModel):
    jour:          str
    heure1:        str
    heure2:        str
    motif:         str
    date_creation: Optional[str] = None   # ← NOUVEAU

class RepriseRequest(BaseModel):
    date_reprise:  str
    adress_medcin: Optional[str] = ""
    date_creation: Optional[str] = None   # ← NOUVEAU

class MaladieRequest(BaseModel):
    date_maladie:  str
    date_reprise:  str
    adress_medcin: Optional[str] = ""
    service:       Optional[str] = ""
    date_creation: Optional[str] = None   # ← NOUVEAU

class FinManquantRequest(BaseModel):
    date_manquant:  str
    heure_manquant: str
    date_creation:  Optional[str] = None  # ← NOUVEAU

class RCRequest(BaseModel):
    service:       str
    nomber_jours:  str
    date_depart:   str
    date_fin:      str
    date_creation: Optional[str] = None   # ← NOUVEAU


# =============================================================================
# ── PE
# =============================================================================

def _build_pe_docx(user, req: PERequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "PE_template.docx"
    if not tp.exists(): raise HTTPException(500, f"Template PE introuvable : {tp}")
    doc = Document(str(tp))
    dash = "—"

    # ← Utilise date_creation manuelle au lieu de date.today()
    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{NOM_PRENOM}}":   user.nom_prenom,
        "{{MATRICULE}}":    user.matricule,
        "{{UNITE}}":        user.unite,
        "{{DESTINATAIRE}}": user.destinataire,
        "{{LIEU}}":         "Casablanca",
        "{{DATE_CREATION}}":date_creation_str,   # ← date manuelle
        "{{HEURE_DEBUT}}":  req.heure_debut if req.type_perm=="Heures" and req.heure_debut else dash,
        "{{HEURE_FIN}}":    req.heure_fin   if req.type_perm=="Heures" and req.heure_fin   else dash,
        "{{DATE_H_J}}":dash,"{{DATE_H_M}}":dash,"{{DATE_H_A}}":dash,
        "{{DEMI1_J}}":dash, "{{DEMI1_M}}":dash, "{{DEMI1_A}}":dash,
        "{{DEMI2_J}}":dash, "{{DEMI2_M}}":dash, "{{DEMI2_A}}":dash,
        "{{JOUR1_J}}":dash, "{{JOUR1_M}}":dash, "{{JOUR1_A}}":dash,
        "{{DEUX1_J}}":dash, "{{DEUX1_M}}":dash, "{{DEUX1_A}}":dash,
        "{{DEUX2_J}}":dash, "{{DEUX2_M}}":dash, "{{DEUX2_A}}":dash,
    }
    if req.type_perm=="Heures" and req.date1:
        j,m,a=split_jma(date.fromisoformat(req.date1)); mapping.update({"{{DATE_H_J}}":j,"{{DATE_H_M}}":m,"{{DATE_H_A}}":a})
    elif req.type_perm=="Demi1" and req.date1:
        j,m,a=split_jma(date.fromisoformat(req.date1)); mapping.update({"{{DEMI1_J}}":j,"{{DEMI1_M}}":m,"{{DEMI1_A}}":a})
    elif req.type_perm=="Demi2" and req.date1:
        j,m,a=split_jma(date.fromisoformat(req.date1)); mapping.update({"{{DEMI2_J}}":j,"{{DEMI2_M}}":m,"{{DEMI2_A}}":a})
    elif req.type_perm=="1Jour" and req.date1:
        j,m,a=split_jma(date.fromisoformat(req.date1)); mapping.update({"{{JOUR1_J}}":j,"{{JOUR1_M}}":m,"{{JOUR1_A}}":a})
    elif req.type_perm=="2Jours" and req.date1 and req.date2:
        j1,m1,a1=split_jma(date.fromisoformat(req.date1)); j2,m2,a2=split_jma(date.fromisoformat(req.date2))
        mapping.update({"{{DEUX1_J}}":j1,"{{DEUX1_M}}":m1,"{{DEUX1_A}}":a1,"{{DEUX2_J}}":j2,"{{DEUX2_M}}":m2,"{{DEUX2_A}}":a2})
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/pe/generate")
def generate_pe(req: PERequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_pe_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "PE", {
        "type_perm":req.type_perm,"date1":req.date1,"date2":req.date2,
        "heure_debut":req.heure_debut,"heure_fin":req.heure_fin,
        "date_creation":req.date_creation,                          # ← sauvegardé en DB
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule,
        "unite":current_user.unite,"destinataire":current_user.destinataire})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"PE_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/pe/regenerate/{doc_id}")
def regenerate_pe(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = PERequest(
        type_perm=meta.get("type_perm","Heures"), date1=meta.get("date1"),
        date2=meta.get("date2"), heure_debut=meta.get("heure_debut"),
        heure_fin=meta.get("heure_fin"),
        date_creation=meta.get("date_creation"),   # ← restauré depuis DB
    )
    docx = _build_pe_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"PE_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── AUTORISATION DE SORTIE
# =============================================================================

def _build_sortie_docx(user, req: SortieRequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "AUTORISATION_SORTIE.docx"
    if not tp.exists(): raise HTTPException(500, "Template AUTORISATION_SORTIE introuvable")
    doc = Document(str(tp))

    # ← Utilise date_creation manuelle
    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{NOM_PRENOM}}":   user.nom_prenom,
        "{{MATRICULE}}":    user.matricule,
        "{{UNITE}}":        user.unite,
        "{{JOUR_COMPLET}}": date.fromisoformat(req.jour).strftime("%d/%m/%Y"),
        "{{HEURE1}}":       fmt_heure(req.heure1),
        "{{HEURE2}}":       fmt_heure(req.heure2),
        "{{MOTIF}}":        req.motif,
        "{{DATE}}":         date_creation_str,   # ← date manuelle
    }
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/sortie/generate")
def generate_sortie(req: SortieRequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_sortie_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "SORTIE", {
        "jour":req.jour,"heure1":req.heure1,"heure2":req.heure2,"motif":req.motif,
        "date_creation":req.date_creation,                          # ← sauvegardé
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule,"unite":current_user.unite})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"SORTIE_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/sortie/regenerate/{doc_id}")
def regenerate_sortie(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = SortieRequest(
        jour=meta["jour"],heure1=meta["heure1"],heure2=meta["heure2"],motif=meta["motif"],
        date_creation=meta.get("date_creation"),   # ← restauré
    )
    docx = _build_sortie_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"SORTIE_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── REPRISE DE SERVICE
# =============================================================================

def _build_reprise_docx(user, req: RepriseRequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "Reprise.docx"
    if not tp.exists(): raise HTTPException(500, "Template Reprise introuvable")
    doc = Document(str(tp))

    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{Nom_Complet}}":   user.nom_prenom,
        "{{Matricule}}":     user.matricule,
        "{{Date_Reprise}}":  date.fromisoformat(req.date_reprise).strftime("%d/%m/%Y"),
        "{{Adress_Medcin}}": req.adress_medcin or "",
        "{{DATE_CREATION}}": date_creation_str,  # ← si présent dans template
        "{{DATE}}":          date_creation_str,  # ← variante
    }
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/reprise/generate")
def generate_reprise(req: RepriseRequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_reprise_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "REPRISE", {
        "date_reprise":req.date_reprise,"adress_medcin":req.adress_medcin,
        "date_creation":req.date_creation,
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"REPRISE_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/reprise/regenerate/{doc_id}")
def regenerate_reprise(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = RepriseRequest(
        date_reprise=meta["date_reprise"],adress_medcin=meta.get("adress_medcin",""),
        date_creation=meta.get("date_creation"),
    )
    docx = _build_reprise_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"REPRISE_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── MALADIE
# =============================================================================

def _build_maladie_docx(user, req: MaladieRequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "MALADIE.docx"
    if not tp.exists(): raise HTTPException(500, "Template MALADIE introuvable")
    doc = Document(str(tp))

    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{nom_complet}}":   user.nom_prenom,
        "{{matricule}}":     user.matricule,
        "{{service}}":       req.service or user.unite,
        "{{date_maladie}}":  date.fromisoformat(req.date_maladie).strftime("%d/%m/%Y"),
        "{{date_reprise}}":  date.fromisoformat(req.date_reprise).strftime("%d/%m/%Y"),
        "{{Adress_Medcin}}": req.adress_medcin or "",
        "{{DATE_CREATION}}": date_creation_str,
        "{{DATE}}":          date_creation_str,
    }
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/maladie/generate")
def generate_maladie(req: MaladieRequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_maladie_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "MALADIE", {
        "date_maladie":req.date_maladie,"date_reprise":req.date_reprise,
        "adress_medcin":req.adress_medcin,"service":req.service,
        "date_creation":req.date_creation,
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"MALADIE_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/maladie/regenerate/{doc_id}")
def regenerate_maladie(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = MaladieRequest(
        date_maladie=meta["date_maladie"],date_reprise=meta["date_reprise"],
        adress_medcin=meta.get("adress_medcin",""),service=meta.get("service",""),
        date_creation=meta.get("date_creation"),
    )
    docx = _build_maladie_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"MALADIE_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── FIN MANQUANT
# =============================================================================

def _build_fin_manquant_docx(user, req: FinManquantRequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "FIN_MANQUANT.docx"
    if not tp.exists(): raise HTTPException(500, "Template FIN_MANQUANT introuvable")
    doc = Document(str(tp))

    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{nom_complet}}":    user.nom_prenom,
        "{{matricule}}":      user.matricule,
        "{{unite}}":          user.unite,
        "{{date_manquant}}":  date.fromisoformat(req.date_manquant).strftime("%d/%m/%Y"),
        "{{heure_manquant}}": req.heure_manquant,
        "{{DATE_CREATION}}":  date_creation_str,
        "{{DATE}}":           date_creation_str,
    }
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/fin-manquant/generate")
def generate_fin_manquant(req: FinManquantRequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_fin_manquant_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "FIN_MANQUANT", {
        "date_manquant":req.date_manquant,"heure_manquant":req.heure_manquant,
        "date_creation":req.date_creation,
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule,"unite":current_user.unite})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"FIN_MANQUANT_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/fin-manquant/regenerate/{doc_id}")
def regenerate_fin_manquant(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = FinManquantRequest(
        date_manquant=meta["date_manquant"],heure_manquant=meta["heure_manquant"],
        date_creation=meta.get("date_creation"),
    )
    docx = _build_fin_manquant_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"FIN_MANQUANT_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── REPOS COMPENSATEUR (RC)
# =============================================================================

def _build_rc_docx(user, req: RCRequest) -> bytes:
    from docx import Document
    tp = TEMPLATES_DIR / "RC.docx"
    if not tp.exists(): raise HTTPException(500, "Template RC introuvable")
    doc = Document(str(tp))

    date_creation_str = parse_date_creation(req.date_creation)

    mapping = {
        "{{nome_complet}}":  user.nom_prenom,
        "{{matricule}}":     user.matricule,
        "{{service}}":       req.service,
        "{{destinataire}}":  user.destinataire,
        "{{nomber_jours}}":  req.nomber_jours,
        "{{date_depart}}":   date.fromisoformat(req.date_depart).strftime("%d/%m/%Y"),
        "{{date_fin}}":      date.fromisoformat(req.date_fin).strftime("%d/%m/%Y"),
        "{{date_creation}}": date_creation_str,   # ← date manuelle
    }
    replace_placeholders(doc, mapping)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

@router.post("/rc/generate")
def generate_rc(req: RCRequest, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    docx = _build_rc_docx(current_user, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    save_doc_record(db, current_user.id, "RC", {
        "service":req.service,"nomber_jours":req.nomber_jours,
        "date_depart":req.date_depart,"date_fin":req.date_fin,
        "date_creation":req.date_creation,
        "nom_prenom":current_user.nom_prenom,"matricule":current_user.matricule})
    safe=re.sub(r"[^A-Za-z0-9_]+","_",current_user.nom_prenom)
    return stream_pdf(pdf, f"RC_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

@router.post("/rc/regenerate/{doc_id}")
def regenerate_rc(doc_id: int, current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    record = db.query(models.DocumentRH).filter(models.DocumentRH.id==doc_id).first()
    if not record: raise HTTPException(404, "Document introuvable")
    original = get_original_user(db, record, current_user)
    meta = record.metadata_doc
    req = RCRequest(
        service=meta["service"],nomber_jours=meta["nomber_jours"],
        date_depart=meta["date_depart"],date_fin=meta["date_fin"],
        date_creation=meta.get("date_creation"),   # ← restauré
    )
    docx = _build_rc_docx(original, req)
    try: pdf = docx_to_pdf_bytes(docx)
    except Exception as e: raise HTTPException(500, f"Conversion PDF échouée : {e}")
    safe=re.sub(r"[^A-Za-z0-9_]+","_",original.nom_prenom)
    return stream_pdf(pdf, f"RC_REGEN_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")


# =============================================================================
# ── HISTORIQUE
# =============================================================================

@router.get("/history")
def get_history(current_user=Depends(auth.get_current_user), db: Session=Depends(get_db)):
    if current_user.role == "admin":
        records = db.query(models.DocumentRH).order_by(models.DocumentRH.created_at.desc()).limit(200).all()
    else:
        records = db.query(models.DocumentRH).filter(
            models.DocumentRH.user_id==current_user.id
        ).order_by(models.DocumentRH.created_at.desc()).all()
    result = []
    for r in records:
        agent = db.query(models.User).filter(models.User.id==r.user_id).first()
        result.append({
            "id": r.id, "type_doc": r.type_doc,
            "created_at": r.created_at.strftime("%d/%m/%Y %H:%M") if r.created_at else "—",
            "agent": agent.nom_prenom if agent else "—",
            "matricule": agent.matricule if agent else "—",
            "metadata": r.metadata_doc,
        })
    return result